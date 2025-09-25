import io
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from django.template import Template as DjangoTemplate, Context


class DocumentGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def generate_pdf(self, template, data):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(template.name, title_style))
        story.append(Spacer(1, 12))
        
        # Process template fields
        fields = json.loads(template.fields) if template.fields else {}
        
        for field_name, field_config in fields.items():
            if field_name in data:
                value = data[field_name]
                
                # Add field label and value
                label_style = ParagraphStyle(
                    'FieldLabel',
                    parent=self.styles['Normal'],
                    fontSize=12,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(f"{field_config.get('label', field_name)}:", label_style))
                story.append(Paragraph(str(value), self.styles['Normal']))
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx(self, template, data):
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(template.name, 0)
        title.alignment = 1  # Center alignment
        
        # Process template fields
        fields = json.loads(template.fields) if template.fields else {}
        
        for field_name, field_config in fields.items():
            if field_name in data:
                value = data[field_name]
                
                # Add field
                p = doc.add_paragraph()
                run = p.add_run(f"{field_config.get('label', field_name)}: ")
                run.bold = True
                p.add_run(str(value))
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_xlsx(self, template, data):
        wb = Workbook()
        ws = wb.active
        if ws is not None:
            ws.title = str(template.name)
            
            # Set title
            title_cell = ws['A1']
            title_cell.value = str(template.name)
            title_cell.font = Font(size=16, bold=True)
            title_cell.alignment = Alignment(horizontal='center')
            
            # Process template fields
            fields = json.loads(template.fields) if template.fields else {}
            row = 3
            
            for field_name, field_config in fields.items():
                if field_name in data:
                    value = data[field_name]
                    
                    # Add field label and value
                    label_cell = ws[f'A{row}']
                    label_cell.value = field_config.get('label', field_name)
                    label_cell.font = Font(bold=True)
                    
                    value_cell = ws[f'B{row}']
                    value_cell.value = str(value)
                    row += 1
        
        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()